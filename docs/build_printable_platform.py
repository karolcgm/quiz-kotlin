from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path


OUT = Path(__file__).with_name("LekcjaLab_opis_platformy_do_druku.docx")

# Resolved preset: narrative_proposal, with a named A4 override for Polish printing.
NAVY = "0B2545"
BLUE = "1F5A85"
CYAN = "2EA8C7"
GOLD = "D99A2B"
INK = "1C2B39"
MUTED = "667585"
LIGHT = "F4F7FA"
PALE_BLUE = "EAF4F8"
PALE_GOLD = "FFF6E6"
WHITE = "FFFFFF"
TABLE_BORDER = "CBD6DF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=TABLE_BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_para(p, before=0, after=8, line=1.25, align=None, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True


def add_text(doc, text, size=11, color=INK, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8, line=1.25, keep=False):
    p = doc.add_paragraph()
    set_para(p, before, after, line, align, keep)
    r = p.add_run(text)
    set_run(r, size, color, bold, italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run(r, 16 if level == 1 else 13 if level == 2 else 12, BLUE if level < 3 else NAVY, True)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Cm(0.95 if level == 0 else 1.4)
    p.paragraph_format.first_line_indent = Cm(-0.48)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_run(r, 10.6)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.95)
    p.paragraph_format.first_line_indent = Cm(-0.48)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_run(r, 10.6)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE, accent=CYAN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    # accent border on left
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    left = borders.find(qn("w:left")) if borders is not None else None
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), accent)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    set_para(p, 0, 3, 1.15)
    r = p.add_run(title)
    set_run(r, 11, NAVY, True)
    p2 = cell.add_paragraph()
    set_para(p2, 0, 0, 1.2)
    r2 = p2.add_run(body)
    set_run(r2, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_band(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    for border in ("top", "left", "bottom", "right"):
        pass
    p = cell.paragraphs[0]
    set_para(p, 0, 0, 1.0)
    r = p.add_run(text.upper())
    set_run(r, 9, WHITE, True)
    r.font.character_spacing = Pt(1)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_two_col_table(doc, rows, header=None, widths=(3000, 6360)):
    table = doc.add_table(rows=1 if header else 0, cols=2)
    if header:
        cells = table.rows[0].cells
        for i, value in enumerate(header):
            set_cell_shading(cells[i], NAVY)
            p = cells[i].paragraphs[0]
            set_para(p, 0, 0, 1.1)
            r = p.add_run(value)
            set_run(r, 10.2, WHITE, True)
    for left, right in rows:
        cells = table.add_row().cells
        for i, value in enumerate((left, right)):
            p = cells[i].paragraphs[0]
            set_para(p, 0, 0, 1.15)
            r = p.add_run(value)
            set_run(r, 10.2, INK, i == 0)
            if i == 0:
                set_cell_shading(cells[i], LIGHT)
    set_table_geometry(table, list(widths))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_three_col_table(doc, rows, header):
    table = doc.add_table(rows=1, cols=3)
    for i, value in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        set_para(p, 0, 0, 1.1)
        r = p.add_run(value)
        set_run(r, 10.1, WHITE, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if i == 0:
                set_cell_shading(cells[i], PALE_BLUE)
            p = cells[i].paragraphs[0]
            set_para(p, 0, 0, 1.12)
            r = p.add_run(value)
            set_run(r, 9.7, INK, i == 0)
    set_table_geometry(table, [2450, 3455, 3455])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run(run, 9, MUTED)


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.8)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for level, size, color, before, after in [(1, 16, BLUE, 18, 10), (2, 13, BLUE, 12, 6), (3, 12, NAVY, 8, 4)]:
        st = styles[f"Heading {level}"]
        st.font.name = "Aptos Display"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]
        st.font.name = "Aptos"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
        st.font.size = Pt(10.6)
        st.font.color.rgb = RGBColor.from_string(INK)

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(hp, 0, 0, 1.0)
    r = hp.add_run("LEKCJALAB  •  MATERIAŁ INFORMACYJNY")
    set_run(r, 8.5, MUTED, True)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(fp, 0, 0, 1.0)
    r = fp.add_run("LekcjaLab  |  opis platformy  •  ")
    set_run(r, 8.5, MUTED)
    add_page_number(fp)


def build():
    doc = Document()
    configure(doc)
    props = doc.core_properties
    props.title = "LekcjaLab — opis platformy do druku"
    props.subject = "Platforma cyfrowego nauczania matematyki"
    props.author = "LekcjaLab"
    props.keywords = "LekcjaLab, matematyka, szkoła, cyfryzacja, pilotaż"

    # First-page centerpiece (proposal_centerpiece with A4 override).
    add_text(doc, "LEKCJALAB", size=11, color=CYAN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, line=1.0)
    add_text(doc, "Cyfrowe środowisko\nprowadzenia lekcji matematyki", size=25, color=NAVY, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.05)
    add_text(doc, "Opis platformy, zgodności programowej i korzyści dla uczniów, nauczycieli oraz szkoły",
             size=13, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=16, line=1.15)
    add_callout(doc, "Jedno zdanie", "LekcjaLab łączy plan nauczania, interaktywną lekcję na tablicy, pracę ucznia na tablecie lub papierze oraz wyniki w jednej mapie umiejętności.", fill=PALE_GOLD, accent=GOLD)
    add_text(doc, "Materiał do rozmowy z dyrekcją, nauczycielami i partnerami wdrożeniowymi", size=10, color=MUTED,
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=4)
    add_text(doc, "Wersja robocza • lipiec 2026", size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()

    add_section_band(doc, "01  Dlaczego LekcjaLab")
    add_heading(doc, "Od katalogu ćwiczeń do spójnej lekcji", 1)
    add_text(doc, "LekcjaLab to platforma edukacyjna zaprojektowana wokół realnego rytmu pracy szkoły. Nauczyciel nie musi składać zajęć z przypadkowych narzędzi: wybiera klasę, dział i temat z planu, uruchamia gotową sekwencję lekcji, a następnie obserwuje, co uczniowie rozumieją i do czego trzeba wrócić. Uczeń otrzymuje krótkie polecenia, modele matematyczne, miejsce na tok rozumowania i natychmiastową informację zwrotną.")
    add_text(doc, "Platforma została pomyślana jako kanał hybrydowy. Może działać z tabletami i tablicą interaktywną, ale podstawowy przebieg lekcji nie powinien zależeć od tego, czy każde urządzenie ma połączenie. Do tych samych umiejętności można wrócić w wersji cyfrowej, na karcie pracy lub w czasie wspólnej rozmowy przy tablicy.")
    add_callout(doc, "Najważniejsza obietnica", "Mniej czasu na organizację i przepisywanie, więcej czasu na rozumowanie, rozmowę o strategiach i świadome poprawianie błędów.")

    add_heading(doc, "Jak działa platforma — od planu do decyzji dydaktycznej", 2)
    for item in [
        "Nauczyciel wybiera szkołę, klasę i temat z wersjonowanego programu.",
        "Otwiera gotowy pakiet: cele, kolejne etapy, przykłady, ćwiczenia i bilet wyjścia.",
        "Uruchamia tablicę oraz opcjonalną sesję live dla tabletów uczniów.",
        "Uczniowie rozwiązują zadania, wpisują odpowiedzi klawiaturą, dotykiem lub na papierze.",
        "Nauczyciel widzi zagregowany stan klasy, typowe błędy i prośby o pomoc — bez publicznego zawstydzania uczniów.",
        "Wyniki trafiają do mapy umiejętności i podpowiadają, co powtórzyć, poprawić albo rozszerzyć."
    ]:
        add_number(doc, item)

    add_section_band(doc, "02  Funkcje platformy")
    add_heading(doc, "Jedno środowisko dla całego obiegu pracy", 1)
    add_two_col_table(doc, [
        ("Program klasy", "Wersjonowana mapa działów, tematów, celów i umiejętności. Nauczyciel widzi miejsce lekcji w planie i postęp realizacji."),
        ("Biblioteka i studio lekcji", "Gotowe pakiety z etapami: start, odkrywanie, ćwiczenie, rozmowa, utrwalenie i bilet wyjścia. Przed zajęciami można wybrać zakres i poziom."),
        ("Tablica + live", "Widok dla całej sali, prywatny pulpit prowadzącego oraz tablet ucznia. Nauczyciel kontroluje tempo, etap, pauzę i moment pokazania rozwiązania."),
        ("Prace i zadania", "Zadania cyfrowe, prace domowe, kartkówki i sprawdziany. Możliwe są terminy, poprawy, zadania dodatkowe i różnicowanie poziomu."),
        ("Druk hybrydowy", "Karty pracy, arkusze A/B, klucze i miejsce na tok rozumowania. Papier jest równoprawnym kanałem, a wynik może trafić do tej samej mapy umiejętności."),
        ("Postępy i dziennik", "Historia prób, wyniki, ślady uczenia, opisowe komentarze i widok tego, co rozumie klasa, a co wymaga kolejnego wyjaśnienia."),
        ("Klasy i zaproszenia", "Nauczyciel tworzy klasy i zaprasza uczniów. Dane są rozdzielone po szkole, więc klasy o tej samej nazwie nie mieszają uczniów."),
        ("Bezpieczeństwo", "Supabase Auth, Postgres i reguły RLS. Rejestracja nauczyciela wymaga ręcznej aktywacji, a uczeń dołącza wyłącznie z ważnego zaproszenia."),
    ], header=("Moduł", "Co daje w praktyce"))

    add_heading(doc, "Matematyka, która wymaga działania", 2)
    add_text(doc, "Modele i zadania nie ograniczają się do odczytania ekranu. Uczeń układa kolejność działań, uzupełnia oś liczbową, wpisuje przeniesienia w dodawaniu i odejmowaniu, buduje piętra mnożenia pisemnego, korzysta z klawiatury numerycznej i uzasadnia wynik. Interakcja ma wspierać myślenie, a nie zastępować zeszyt.")
    add_bullet(doc, "W zadaniach pisemnych można uzupełniać cyfry krok po kroku, zamiast otrzymywać wyłącznie pole na końcowy wynik.")
    add_bullet(doc, "W przykładach pojawiają się liczby dwu- i trzycyfrowe, potęgi, szacowanie oraz proste zadania tekstowe.")
    add_bullet(doc, "Wariant papierowy zachowuje miejsce na zapis strategii, odpowiedź słowną i pracę bez urządzenia.")

    add_section_band(doc, "03  Zgodność z kształceniem")
    add_heading(doc, "Platforma wspiera realizację programu, nie zastępuje nauczyciela", 1)
    add_text(doc, "LekcjaLab jest osadzona w wersjonowanym planie matematyki dla klasy V. Tematy mają cele, umiejętności, kolejność i status (obowiązkowy, zalecany, opcjonalny lub rozszerzający). Dzięki temu szkoła może prowadzić własny plan, a jednocześnie zachować czytelną mapę tego, czego uczniowie się uczą.")
    add_callout(doc, "Ważne zastrzeżenie", "Platforma jest narzędziem wspierającym pracę szkoły. Nie jest deklaracją formalnej akredytacji ani gotowym programem nauczania zatwierdzonym przez MEN. Ostateczne dopasowanie do programu i zasad oceniania należy do szkoły i nauczyciela.", fill=PALE_GOLD, accent=GOLD)
    add_heading(doc, "Przykład: dział „Liczby i działania”", 2)
    add_two_col_table(doc, [
        ("Zapisywanie i porównywanie liczb", "Oś liczbowa, liczby w setkach, klawiatura cyfr i porównywanie wielkości."),
        ("Dodawanie i odejmowanie w pamięci", "Proste przypadki na liczbach dwu- i trzycyfrowych oraz utrwalanie nazw składników, sumy, odjemnej i odjemnika."),
        ("Mnożenie i dzielenie w pamięci", "Strategie rachunkowe, przykłady z potęgami, w tym 30²."),
        ("Kolejność działań", "Kalkulatorowa klawiatura, małe liczby, nawiasy i potęgi, np. (5−3)² × 4²."),
        ("Szacowanie", "Zaokrąglanie liczb i ocenianie rzędu wielkości wyniku w dodawaniu, odejmowaniu, mnożeniu i dzieleniu."),
        ("Działania pisemne", "Dodawanie, odejmowanie z przeniesieniem oraz mnożenie pisemne z uzupełnianiem kolejnych pięter i zer."),
    ], header=("Obszar", "Przykładowe doświadczenie ucznia"))
    add_text(doc, "Każdy temat może zostać poprowadzony cyfrowo, na tablicy lub jako wydruk. W przyszłości ten sam kontrakt programu może obsłużyć kolejne klasy i przedmioty bez kopiowania całej platformy.")

    add_section_band(doc, "04  Uczeń i motywacja")
    add_heading(doc, "Uczeń chce wracać, bo widzi sens i postęp", 1)
    add_text(doc, "Uczeń niechętnie sięga po szary zeszyt, w którym widzi tylko tekst, poprawki i często nieczytelne notatki. W LekcjaLab ponownie wykonuje działanie: wybiera, wpisuje, porównuje, sprawdza i poprawia. Ekran nie ma zastąpić zapisu — ma zaprosić do aktywnego myślenia, a potem pomóc uporządkować rozwiązanie.")
    add_bullet(doc, "Każda lekcja ma małe kroki i jasny cel, więc uczeń wie, co właśnie ćwiczy.")
    add_bullet(doc, "Natychmiastowa informacja zwrotna pokazuje, co było poprawne i nad czym trzeba popracować.")
    add_bullet(doc, "Uczeń może wrócić do lekcji w domu, utrwalić materiał i zobaczyć własną mapę postępów.")
    add_bullet(doc, "Za ukończone lekcje, poprawy i zadania dodatkowe zdobywa unikalne cyfrowe nagrody budujące prywatny cyfrowy dorobek.")
    add_bullet(doc, "Nagrody nie muszą oznaczać publicznego rankingu — ważniejsze jest poczucie sprawczości, kolekcjonowania i rozwoju.")

    add_heading(doc, "Gry matematyczne i praca zespołowa", 2)
    add_text(doc, "Kierunek rozwoju platformy obejmuje gry matematyczne do prezentacji w sali oraz tryby rywalizacji i współpracy między zespołami. Klasa może rozwiązywać zadania jako drużyny matematyczne, prowadzić burze mózgów, wybierać strategie i uczestniczyć w konkursach. Zwycięstwo nie powinno zależeć wyłącznie od szybkości — można punktować poprawność, uzasadnienie i pomoc kolegom.")
    add_callout(doc, "Mechanizm przyciągania", "LekcjaLab łączy ciekawość gry z konkretną umiejętnością: uczeń zdobywa nagrodę za wykonane działanie, a nauczyciel nadal widzi, czego uczeń się nauczył.")

    add_section_band(doc, "05  Korzyści")
    add_heading(doc, "Korzyści dla trzech stron", 1)
    add_three_col_table(doc, [
        ("Uczeń", "Aktywne zadania zamiast biernego czytania; bezpieczne poprawianie; widoczny postęp; powtórka w domu; nagrody i własny cyfrowy dorobek.", "Większa samodzielność i odwaga w podejmowaniu kolejnych prób."),
        ("Nauczyciel", "Gotowa struktura lekcji; szybkie zadania i wydruki; poprawy i prace dodatkowe; opisowa ocena; feedback uczniów; statystyki do planowania kolejnych zajęć.", "Mniej pracy administracyjnej, szybsza reakcja na braki i lepsze różnicowanie."),
        ("Szkoła", "Spójny standard pracy na tabletach i tablicy; dane rozdzielone po szkole; kanał cyfrowy i papierowy; pilotaż możliwy bez opłat licencyjnych.", "Praktyczny element planu cyfryzacji i materiał do mierzalnej innowacji pedagogicznej."),
    ], header=("Odbiorca", "Co otrzymuje", "Efekt"))

    add_heading(doc, "Informacja zwrotna zamiast samej oceny", 2)
    add_text(doc, "Uczeń może ocenić lekcję i zaznaczyć, co rozumie, a czego nie. Nauczyciel otrzymuje nie tylko wynik, lecz także sygnał: „umiem samodzielnie”, „potrzebuję przykładu” albo „nie rozumiem pojęcia”. Agregowane statystyki pokazują, gdzie klasa traci punkty i jakie błędy powtarzają się najczęściej. Pozwala to poprawić program lub tempo zanim braki staną się uciążliwe.")
    add_text(doc, "Opisowa ocena może towarzyszyć punktom: nauczyciel zapisuje krótką informację o strategii, postępie i następnym kroku. Uczeń wie wtedy nie tylko, ile zdobył, ale także co zrobić dalej.")

    add_section_band(doc, "06  Cyfryzacja i wdrożenie")
    add_heading(doc, "Naturalne uzupełnienie szkolnej infrastruktury", 1)
    add_text(doc, "Jeżeli szkoła ma tablety i tablicę interaktywną, LekcjaLab porządkuje ich użycie wokół celu lekcji. Tablica służy do wspólnego modelu i rozmowy, tablet do indywidualnej próby, a papier pozostaje bezpiecznym planem B. Platforma nie wymaga, aby technologia była celem samym w sobie.")
    add_bullet(doc, "Wspiera kompetencje cyfrowe: bezpieczne logowanie, pracę z informacją, korzystanie z narzędzia i odpowiedzialne przesyłanie pracy.")
    add_bullet(doc, "Łączy cyfrowe i papierowe dowody uczenia w jednym obiegu, zamiast tworzyć dwa niezależne dzienniki.")
    add_bullet(doc, "Może dostarczyć szkole mierzalnych danych do oceny innowacji: aktywność, ukończenia, poprawy, wyniki i opinie uczestników.")
    add_bullet(doc, "Uwzględnia prywatność dzieci: brak reklam i publicznych rankingów z nazwiskami, separacja szkół i ograniczenie danych do potrzeb edukacyjnych.")

    add_heading(doc, "Rekomendowany pilotaż w jednej szkole", 2)
    add_text(doc, "Najbezpieczniejszym sposobem rozpoczęcia jest pilotaż w jednej klasie, w której nauczyciel zna realia szkoły i może zebrać opinie. W opisanej sytuacji naturalnym liderem może być nauczycielka pracująca w szkole wyposażonej w tablety i tablicę interaktywną. Ponieważ aplikacja jest obecnie darmowa, szkoła może skupić się na wartości dydaktycznej, a nie na decyzji zakupowej.")
    for item in [
        "Ustalić z dyrektorem zakres: jedna klasa, jeden dział i 6–8 tygodni pracy.",
        "Przeprowadzić krótką diagnozę przed i po cyklu oraz zachować porównywalne prace papierowe/cyfrowe.",
        "Zebrać osobne opinie uczniów, nauczycieli, dyrektora i rodziców.",
        "Udokumentować przykłady: czas uruchomienia lekcji, liczba popraw, najczęstsze trudności i zaangażowanie uczniów.",
        "Po pilotażu zdecydować, które tematy i tryby warto rozszerzyć na kolejne klasy."
    ]:
        add_number(doc, item)
    add_callout(doc, "Warunek odpowiedzialnego wdrożenia", "Przed uruchomieniem szkoła powinna uzgodnić zasady dostępu, informację dla rodziców, retencję danych i sposób wykorzystania wyników. Technologia ma wzmacniać relację nauczyciel–uczeń, a nie zastępować jej.", fill=PALE_GOLD, accent=GOLD)

    add_section_band(doc, "07  Podsumowanie")
    add_heading(doc, "LekcjaLab w jednym akapicie", 1)
    add_text(doc, "LekcjaLab to spójne centrum prowadzenia matematyki: od planu klasy, przez interaktywną lekcję na tablicy i tabletach, po zadanie, wydruk, poprawę, opisową informację zwrotną i mapę umiejętności. Uczeń może ćwiczyć w szkole i w domu, zdobywać cyfrowe nagrody oraz wracać do treści, których jeszcze nie opanował. Nauczyciel zyskuje czas, widoczność postępów i narzędzia do różnicowania. Szkoła otrzymuje praktyczny, możliwy do zmierzenia element cyfryzacji i innowacji pedagogicznej.")
    add_text(doc, "Najlepszym następnym krokiem nie jest wielkie wdrożenie, lecz dobrze opisany pilotaż: jedna szkoła, jedna klasa, konkretne cele, opinie wszystkich uczestników i decyzja oparta na dowodach.", size=11.5, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=10, line=1.2)
    add_text(doc, "Kontakt / notatki do rozmowy wdrożeniowej: ________________________________________________\n__________________________________________________________________________________________", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.3)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
